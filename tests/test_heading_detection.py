"""Detecting headings that are faked with paragraph formatting (no Heading style).

Score/threshold/weights live in a configurable HeadingConfig; definitive signals
(w:outlineLvl, Title style) bypass scoring; a date/time line is gated out.
"""

from __future__ import annotations

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from parsers.base import HeadingBlock, ParagraphBlock
from parsers.docx_parser import DocxParser, HeadingConfig

_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _doc(build) -> "docx.Document":
    d = docx.Document()
    build(d)
    return d


def _parse(tmp_path, d, cfg=None):
    p = tmp_path / "h.docx"
    d.save(str(p))
    return DocxParser(heading_config=cfg).parse(p, "h")


def _heading(doc, text):
    return next((b for b in doc.blocks
                 if isinstance(b, HeadingBlock) and b.text == text), None)


# --- scored promotion -------------------------------------------------------


def test_bold_caps_short_paragraph_becomes_heading(tmp_path):
    d = docx.Document()
    d.add_paragraph().add_run("OVERVIEW").bold = True      # bold+caps+vshort+no_punct
    d.add_paragraph("Some ordinary body text follows the heading here.")
    doc = _parse(tmp_path, d)
    h = _heading(doc, "OVERVIEW")
    assert h is not None and h.level == 2                  # default level H2


def test_ordinary_sentence_is_not_a_heading(tmp_path):
    d = docx.Document()
    d.add_paragraph("This is a normal sentence of body text that ends properly.")
    d.add_paragraph("Another paragraph.")
    doc = _parse(tmp_path, d)
    assert not any(isinstance(b, HeadingBlock) for b in doc.blocks)


# --- configurability --------------------------------------------------------


def test_threshold_is_configurable(tmp_path):
    d = docx.Document()
    d.add_paragraph().add_run("OVERVIEW").bold = True
    d.add_paragraph("Body text after the heading.")
    # An unreachable threshold must suppress the very same promotion.
    doc = _parse(tmp_path, d, cfg=HeadingConfig(threshold=99))
    assert _heading(doc, "OVERVIEW") is None
    assert any(isinstance(b, ParagraphBlock) and b.text == "OVERVIEW"
               for b in doc.blocks)


def test_weight_zero_disables_a_clue(tmp_path):
    d = docx.Document()
    d.add_paragraph().add_run("OVERVIEW").bold = True
    d.add_paragraph("Body text after the heading.")
    # Kill the bold weight -> OVERVIEW (caps+vshort+no_punct = 6) drops below 7.
    doc = _parse(tmp_path, d, cfg=HeadingConfig(bold=0.0))
    assert _heading(doc, "OVERVIEW") is None


# --- gates ------------------------------------------------------------------


def test_date_line_is_gated_out(tmp_path):
    d = docx.Document()
    d.add_paragraph().add_run("MARCH 31, 2014").bold = True   # bold caps but a date
    d.add_paragraph("Body text after the date line.")
    doc = _parse(tmp_path, d)
    assert _heading(doc, "MARCH 31, 2014") is None


def test_date_gate_can_be_disabled(tmp_path):
    d = docx.Document()
    d.add_paragraph().add_run("MARCH 31, 2014").bold = True
    d.add_paragraph("Body text after the date line.")
    doc = _parse(tmp_path, d, cfg=HeadingConfig(enable_date_gate=False))
    assert _heading(doc, "MARCH 31, 2014") is not None


# --- definitive signals + levels --------------------------------------------


def test_outline_level_promotes_and_sets_level(tmp_path):
    d = docx.Document()
    p = d.add_paragraph("subtle heading with no other clue at all here maybe")
    p._p.get_or_add_pPr().append(
        parse_xml(f'<w:outlineLvl {_W} w:val="1"/>'))     # val 1 -> level 2
    d.add_paragraph("body")
    doc = _parse(tmp_path, d)
    h = next(b for b in doc.blocks if isinstance(b, HeadingBlock))
    assert h.level == 2


def test_numbering_depth_sets_level(tmp_path):
    d = docx.Document()
    r = d.add_paragraph().add_run("1.2 Methods and Materials")
    r.bold = True
    d.add_paragraph("body text")
    doc = _parse(tmp_path, d)
    h = _heading(doc, "1.2 Methods and Materials")
    assert h is not None and h.level == 2                 # depth of "1.2"


def test_larger_font_promotes_to_h1(tmp_path):
    from docx.shared import Pt
    d = docx.Document()
    d.styles["Normal"].font.size = Pt(11)                 # baseline 22 half-points
    run = d.add_paragraph().add_run("Annual Report")
    run.bold = True
    run.font.size = Pt(20)                                # ~1.8x -> xlarge -> H1
    d.add_paragraph("Body text of the report.")
    doc = _parse(tmp_path, d)
    h = _heading(doc, "Annual Report")
    assert h is not None and h.level == 1
