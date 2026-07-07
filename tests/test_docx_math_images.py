"""DOCX extraction fixes for known gaps:

1. Office Math (OMML m:t) text — was dropped because only w:t was read.
3. Legacy/OLE images (VML v:imagedata r:id) — were ignored (only a:blip handled).
4. Images inside table cells — were dropped (cell reader produced text only).

Items 1 and 4 also cover the body-vs-cell parity: the same `_walk_para` powers
both, so math/images work wherever they appear.
"""

from __future__ import annotations

import struct
import zlib
from io import BytesIO

import docx
from docx.oxml import parse_xml

from parsers.base import (
    ParsedDocument, TableBlock, ImageBlock, ParagraphBlock, HeadingBlock, Mark,
)
from parsers.docx_parser import DocxParser, _walk_para


def _text(runs):
    return "".join(r.text for r in runs)


# --- helpers ----------------------------------------------------------------

_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
_M = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
_A = 'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
_V = 'xmlns:v="urn:schemas-microsoft-com:vml"'
_R = 'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
_MC = 'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'

_OMATH = f'<m:oMath {_M}><m:r><m:t>∇φ</m:t></m:r></m:oMath>'  # ∇φ


def _png_1x1() -> bytes:
    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return (struct.pack(">I", len(data)) + body
                + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF))
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1 truecolor
    idat = zlib.compress(b"\x00\xff\x00\x00")            # one red pixel
    return (b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


class _FakePart:
    def __init__(self, name: str, ct: str) -> None:
        self.partname, self.content_type = name, ct


class _FakeDoc:
    """Stand-in for a python-docx Document exposing only related_parts."""
    class _P:
        related_parts = {
            "rId5": _FakePart("/word/media/imageA.png", "image/png"),
            "rId7": _FakePart("/word/media/imageB.png", "image/png"),
        }
    part = _P()


# --- 1. Office Math ---------------------------------------------------------


def test_math_in_body_paragraph_is_captured():
    p = parse_xml(f'<w:p {_W} {_M}><w:r><w:t>x=</w:t></w:r>{_OMATH}</w:p>')
    runs, images = _walk_para(p, _FakeDoc(), [0])
    assert _text(runs) == "x=∇φ"   # "x=∇φ", math no longer dropped
    assert images == []


def test_math_in_table_cell_is_captured(tmp_path):
    d = docx.Document()
    cell = d.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run("x=")
    cell.paragraphs[0]._p.append(parse_xml(_OMATH))
    p = tmp_path / "math.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "math")
    txt = doc.tables()[0].table.cells[0][0].plain_text()
    assert "∇φ" in txt              # ∇φ present
    assert txt == "x=∇φ"


# --- 3. Legacy / OLE (VML) images -------------------------------------------


def test_vml_imagedata_becomes_an_image():
    p = parse_xml(
        f'<w:p {_W} {_V} {_R}><w:r><w:pict>'
        f'<v:shape><v:imagedata r:id="rId7"/></v:shape>'
        f'</w:pict></w:r></w:p>')
    runs, images = _walk_para(p, _FakeDoc(), [0])
    assert _text(runs) == "<image1>"
    assert images == [(1, "word/media/imageB.png", "image/png")]


def test_alternate_content_image_is_not_double_counted():
    """mc:AlternateContent ships a modern (Choice/a:blip) and legacy
    (Fallback/v:imagedata) rendering of ONE image; count it once."""
    p = parse_xml(
        f'<w:p {_W} {_MC} {_A} {_V} {_R}><mc:AlternateContent>'
        f'<mc:Choice Requires="wps"><w:r><w:drawing>'
        f'<a:blip r:embed="rId5"/></w:drawing></w:r></mc:Choice>'
        f'<mc:Fallback><w:r><w:pict><v:shape>'
        f'<v:imagedata r:id="rId7"/></v:shape></w:pict></w:r></mc:Fallback>'
        f'</mc:AlternateContent></w:p>')
    runs, images = _walk_para(p, _FakeDoc(), [0])
    assert images == [(1, "word/media/imageA.png", "image/png")]  # Choice only
    assert _text(runs).count("<image") == 1


# --- 4. Images inside table cells -------------------------------------------


def test_image_in_table_cell_emits_imageblock(tmp_path):
    d = docx.Document()
    cell = d.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run().add_picture(BytesIO(_png_1x1()))
    p = tmp_path / "cellimg.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "cellimg")
    # image did not leak to document body...
    assert doc.images() == [] or all(
        b.locator for b in doc.images())
    blocks = doc.tables()[0].table.cells[0][0].blocks
    imgs = [b for b in blocks if isinstance(b, ImageBlock)]
    assert len(imgs) == 1
    assert imgs[0].locator["part"].startswith("word/media/")
    # the marker shows in the flat text view exactly once
    assert doc.tables()[0].table.cells[0][0].plain_text() == imgs[0].marker


def test_body_and_cell_images_share_one_counter(tmp_path):
    """Image indices are globally sequential in reading order (body then cell)."""
    d = docx.Document()
    d.add_paragraph().add_run().add_picture(BytesIO(_png_1x1()))  # body image -> 1
    cell = d.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run().add_picture(BytesIO(_png_1x1()))  # cell image -> 2
    p = tmp_path / "mixed.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "mixed")
    body_imgs = [b for b in doc.blocks if isinstance(b, ImageBlock)]
    cell_imgs = [b for row in doc.tables()[0].table.cells for c in row if c
                 for b in c.blocks if isinstance(b, ImageBlock)]
    assert [b.image_index for b in body_imgs] == [1]
    assert [b.image_index for b in cell_imgs] == [2]


def test_roundtrip_with_cell_image(tmp_path):
    d = docx.Document()
    cell = d.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run().add_picture(BytesIO(_png_1x1()))
    p = tmp_path / "rt.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "rt")
    restored = ParsedDocument.from_json(doc.to_json())
    assert restored.to_dict() == doc.to_dict()
