"""Nested tables inside DOCX table cells (IR v2).

A Word table cell may contain another table. The v2 IR models a cell as a list of
blocks (`Cell.blocks`), so a nested table is preserved as a real `TableBlock`
inside the parent cell instead of being flattened into unusable concatenated text
like "AdTutarSeyahat695.50".
"""

from __future__ import annotations

import docx
import pytest

from parsers.base import (
    ParsedDocument, TableBlock, ParagraphBlock, Cell, TableData, text_cell,
)
from parsers.docx_parser import DocxParser, _MAX_TABLE_DEPTH


# --- fixtures ---------------------------------------------------------------


def _finance_docx(path) -> None:
    """Outer 1x1 table; its cell holds a 'Departman: Finans' paragraph followed
    by a 2x2 inner table (Ad/Tutar, Seyahat/695.50)."""
    d = docx.Document()
    outer = d.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    cell.paragraphs[0].text = "Departman: Finans"
    inner = cell.add_table(rows=2, cols=2)
    inner.cell(0, 0).text = "Ad"
    inner.cell(0, 1).text = "Tutar"
    inner.cell(1, 0).text = "Seyahat"
    inner.cell(1, 1).text = "695.50"
    d.save(str(path))


def _deeply_nested_docx(path, depth: int) -> None:
    """A chain of `depth` tables, each nested in the single cell of its parent,
    innermost cell holding the text 'DEEP'."""
    d = docx.Document()
    outer = d.add_table(rows=1, cols=1)
    cur = outer.cell(0, 0)
    for _ in range(depth - 1):
        cur = cur.add_table(rows=1, cols=1).cell(0, 0)
    cur.paragraphs[0].text = "DEEP"
    d.save(str(path))


@pytest.fixture()
def finance_doc(tmp_path) -> ParsedDocument:
    p = tmp_path / "finance.docx"
    _finance_docx(p)
    return DocxParser().parse(p, "finance")


# --- structure --------------------------------------------------------------


def test_inner_table_is_a_real_block_inside_the_cell(finance_doc):
    outer = finance_doc.tables()[0]
    cell = outer.table.cells[0][0]
    assert isinstance(cell, Cell)
    kinds = [type(b).__name__ for b in cell.blocks]
    assert kinds == ["ParagraphBlock", "TableBlock"]

    inner = next(b for b in cell.blocks if isinstance(b, TableBlock))
    assert inner.table.n_rows == 2 and inner.table.n_cols == 2
    assert inner.table.cells[0][0].plain_text() == "Ad"
    assert inner.table.cells[1][1].plain_text() == "695.50"


def test_plain_text_is_structured_never_concatenated_garbage(finance_doc):
    cell = finance_doc.tables()[0].table.cells[0][0]
    text = cell.plain_text()
    assert "Departman: Finans" in text
    assert "Seyahat" in text
    assert "695.50" in text
    # the exact bug we set out to kill:
    assert "AdTutarSeyahat695.50" not in text
    # Seyahat and its amount must be separated, not fused:
    assert "Seyahat695.50" not in text


# --- round-trip / losslessness ---------------------------------------------


def test_ir_roundtrips_through_json(finance_doc):
    restored = ParsedDocument.from_json(finance_doc.to_json())
    assert restored.to_dict() == finance_doc.to_dict()


def test_nested_structure_survives_roundtrip(finance_doc):
    restored = ParsedDocument.from_json(finance_doc.to_json())
    cell = restored.tables()[0].table.cells[0][0]
    inner = [b for b in cell.blocks if isinstance(b, TableBlock)]
    assert len(inner) == 1
    assert inner[0].table.cells[1][0].plain_text() == "Seyahat"


def test_legacy_v1_string_cell_deserializes(finance_doc):
    """Old (v1) IR stored a cell as a bare string; readers must still accept it."""
    data = finance_doc.to_dict()
    # forge a v1-style table: cells are plain strings / null
    data["blocks"].append({
        "type": "table", "id": "legacy",
        "table": {"n_rows": 1, "n_cols": 2,
                  "cells": [["Ad", "Tutar"]], "merges": []},
    })
    restored = ParsedDocument.from_dict(data)
    legacy = restored.block_by_id("legacy")
    assert isinstance(legacy.table.cells[0][0], Cell)
    assert legacy.table.cells[0][0].plain_text() == "Ad"


@pytest.mark.skip(reason="lossless-edit / DOCX anchor-writer subsystem does not "
                         "exist yet; losslessness is covered here by IR round-trip "
                         "(test_ir_roundtrips_through_json). Wire this up when the "
                         "write-back module lands.")
def test_lossless_edit_writeback_preserves_nested_table():
    ...


# --- depth guard ------------------------------------------------------------


def test_deeply_nested_table_does_not_crash(tmp_path):
    p = tmp_path / "deep.docx"
    _deeply_nested_docx(p, depth=12)  # deeper than _MAX_TABLE_DEPTH
    doc = DocxParser().parse(p, "deep")  # must not raise
    assert len(doc.tables()) == 1
    # content past the guard is degraded to text but never lost:
    assert "DEEP" in doc.to_json()


def test_depth_guard_constant_is_sane():
    assert 1 <= _MAX_TABLE_DEPTH <= 100


# --- text-cell helper -------------------------------------------------------


def test_text_cell_wraps_and_preserves_none():
    assert text_cell(None) is None                     # merge-covered / absent
    assert text_cell("").blocks == []                  # empty cell, no paragraph
    c = text_cell("hi")
    assert isinstance(c.blocks[0], ParagraphBlock)
    assert c.plain_text() == "hi"
