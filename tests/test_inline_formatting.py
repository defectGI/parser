"""Inline text formatting (bold/italic/underline/strike/super/subscript).

The IR keeps `text` as the canonical plain view and adds an optional parallel
`runs` list carrying the semantic marks. `"".join(r.text for r in runs) == text`.
DOCX populates it from w:rPr; other parsers do not yet.
"""

from __future__ import annotations

import docx
from docx.oxml import parse_xml

from parsers.base import (
    ParsedDocument, ParagraphBlock, Mark, InlineRun, finalize_runs,
)
from parsers.docx_parser import DocxParser, _walk_para

_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _para(inner_xml: str):
    return parse_xml(f"<w:p {_W}>{inner_xml}</w:p>")


class _FakeDoc:
    class _P:
        related_parts: dict = {}
    part = _P()


def _run(text: str, *rpr: str) -> str:
    props = f"<w:rPr>{''.join(rpr)}</w:rPr>" if rpr else ""
    return f"<w:r>{props}<w:t>{text}</w:t></w:r>"


# --- run extraction ---------------------------------------------------------


def test_bold_and_italic_runs():
    p = _para(_run("plain ") + _run("bold", "<w:b/>") + _run(" and ")
              + _run("ital", "<w:i/>"))
    runs, _ = _walk_para(p, _FakeDoc(), [0])
    assert "".join(r.text for r in runs) == "plain bold and ital"
    marked = {r.text: r.marks for r in runs}
    assert marked["bold"] == (Mark.BOLD,)
    assert marked["ital"] == (Mark.ITALIC,)
    assert marked["plain "] == ()


def test_superscript_and_subscript_preserved():
    # "x2" where the 2 is a superscript, and "H2O" where the 2 is subscript.
    p = _para(_run("x") + _run("2", '<w:vertAlign w:val="superscript"/>')
              + _run(" H") + _run("2", '<w:vertAlign w:val="subscript"/>')
              + _run("O"))
    runs, _ = _walk_para(p, _FakeDoc(), [0])
    by = [(r.text, r.marks) for r in runs]
    assert ("2", (Mark.SUPERSCRIPT,)) in by
    assert ("2", (Mark.SUBSCRIPT,)) in by
    # the plain text alone is ambiguous (x2 / H2O) — the marks disambiguate it:
    assert "".join(r.text for r in runs) == "x2 H2O"


def test_combined_marks_on_one_run():
    p = _para(_run("both", "<w:b/>", "<w:i/>",
                   '<w:vertAlign w:val="superscript"/>'))
    runs, _ = _walk_para(p, _FakeDoc(), [0])
    assert set(runs[0].marks) == {Mark.BOLD, Mark.ITALIC, Mark.SUPERSCRIPT}


def test_explicitly_disabled_toggle_is_not_a_mark():
    p = _para(_run("off", '<w:b w:val="0"/>'))
    runs, _ = _walk_para(p, _FakeDoc(), [0])
    assert runs[0].marks == ()


def test_adjacent_same_mark_runs_merge():
    p = _para(_run("Hel", "<w:b/>") + _run("lo", "<w:b/>"))
    runs, _ = _walk_para(p, _FakeDoc(), [0])
    assert len(runs) == 1
    assert runs[0].text == "Hello" and runs[0].marks == (Mark.BOLD,)


# --- block wiring + serialization -------------------------------------------


def test_paragraph_block_carries_runs(tmp_path):
    d = docx.Document()
    par = d.add_paragraph()
    par.add_run("normal ")
    par.add_run("bold").bold = True
    p = tmp_path / "fmt.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "fmt")
    blk = next(b for b in doc.blocks if isinstance(b, ParagraphBlock))
    assert blk.text == "normal bold"                     # plain view intact
    assert "".join(r.text for r in blk.runs) == blk.text  # runs reconstruct text
    assert any(r.marks == (Mark.BOLD,) and r.text == "bold" for r in blk.runs)


def test_unformatted_paragraph_stores_no_runs(tmp_path):
    d = docx.Document()
    d.add_paragraph("just plain text")
    p = tmp_path / "plain.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "plain")
    blk = next(b for b in doc.blocks if isinstance(b, ParagraphBlock))
    assert blk.text == "just plain text"
    assert blk.runs == []                    # lean: no redundant runs, no JSON bloat
    assert "runs" not in blk.to_dict()


def test_runs_survive_json_roundtrip(tmp_path):
    d = docx.Document()
    par = d.add_paragraph()
    par.add_run("a")
    par.add_run("b").italic = True
    p = tmp_path / "rt.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "rt")
    restored = ParsedDocument.from_json(doc.to_json())
    assert restored.to_dict() == doc.to_dict()
    blk = next(b for b in restored.blocks if isinstance(b, ParagraphBlock))
    assert any(Mark.ITALIC in r.marks for r in blk.runs)


def test_formatting_in_table_cell(tmp_path):
    d = docx.Document()
    cell = d.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run("bold cell").bold = True
    p = tmp_path / "cellfmt.docx"
    d.save(str(p))

    doc = DocxParser().parse(p, "cellfmt")
    cell_blocks = doc.tables()[0].table.cells[0][0].blocks
    par = next(b for b in cell_blocks if isinstance(b, ParagraphBlock))
    assert par.text == "bold cell"
    assert any(Mark.BOLD in r.marks for r in par.runs)


# --- finalize_runs helper ---------------------------------------------------


def test_finalize_strips_edges_and_drops_empty():
    runs = finalize_runs([
        InlineRun("  lead", ()), InlineRun("", (Mark.BOLD,)),
        InlineRun(" mid ", ()), InlineRun("tail  ", ()),
    ])
    assert "".join(r.text for r in runs) == "lead mid tail"
